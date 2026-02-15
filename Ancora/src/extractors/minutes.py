"""Extract action items and decisions from Ancora meeting minutes."""
import os
import re
import glob
from docx import Document
from src.config import DATA_MINUTES


def _parse_date_from_filename(filename):
    """Extract date from filename like '02102026 Ancora Meeting Minutes...'."""
    # Try MMDDYYYY format
    match = re.match(r'(\d{2})(\d{2})(\d{4})', filename)
    if match:
        month, day, year = match.groups()
        return f"{year}-{month}-{day}"
    # Try YYYY_MM_DD format
    match = re.search(r'(\d{4})_(\d{2})_(\d{2})', filename)
    if match:
        return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
    return None


def _parse_minutes_docx(filepath):
    """Parse a meeting minutes DOCX and extract structured info."""
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract table data
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)

    result = {
        "summary": "",
        "key_metrics": [],
        "action_items": [],
        "positive_outlook": [],
        "negative_outlook": [],
        "full_text": "\n".join(paragraphs),
    }

    section = None
    for para in paragraphs:
        lower = para.lower()

        # Detect sections (Chinese headers used in Ancora minutes)
        if "简介" in para:
            section = "intro"
        elif "现状" in para or "运营状态" in para:
            section = "status"
        elif "最新进展" in para:
            section = "updates"
        elif "积极" in para and ("因素" in para or "面" in para):
            section = "positive"
        elif "消极" in para and ("因素" in para or "面" in para):
            section = "negative"
        elif "投资" in para and "决策" in para:
            section = "decisions"
        elif "关注" in para:
            section = "watchlist"
        elif "展望" in para:
            section = "outlook"

        # Extract based on section
        if section == "intro" and para != "简介" and not result["summary"]:
            result["summary"] = para

        elif section == "positive":
            if "：" in para or ":" in para:
                result["positive_outlook"].append(para)

        elif section == "negative":
            if "：" in para or ":" in para:
                result["negative_outlook"].append(para)

        elif section == "decisions":
            if any(kw in para for kw in ["批准", "确认", "决定", "需要", "建立",
                                          "审核", "监督", "评估", "动态", "输出"]):
                responsible = "Unknown"
                if "Steven" in para or "steven" in lower:
                    responsible = "Steven (Pondmoon)"
                elif "Nicole" in para or "nicole" in lower:
                    responsible = "Nicole (Pondmoon)"
                elif "Patrick" in para or "patrick" in lower:
                    responsible = "Patrick (Pondmoon)"
                elif "Victoria" in para or "victoria" in lower:
                    responsible = "Victoria (Greystar)"
                elif "Greystar" in para or "greystar" in lower:
                    responsible = "Greystar"
                elif "Pondmoon" in para:
                    responsible = "Pondmoon"

                result["action_items"].append({
                    "action": para.strip("- ").strip(),
                    "responsible": responsible,
                    "status": "pending",
                    "category": "decision",
                })

        elif section == "watchlist":
            if "：" in para or ":" in para:
                result["action_items"].append({
                    "action": f"[Watch] {para.strip()}",
                    "responsible": "",
                    "status": "monitoring",
                    "category": "watchlist",
                })

    # Extract key metrics from table
    for row in table_data:
        if len(row) >= 2 and row[0] and row[1]:
            result["key_metrics"].append({
                "metric": row[0],
                "value": row[1],
                "notes": row[2] if len(row) > 2 else "",
            })

    return result


def extract():
    """Extract action items from all meeting minutes files.

    Returns a list of action entries for the actions log.
    """
    actions = []

    if not os.path.exists(DATA_MINUTES):
        print("  [minutes] Data_Minutes directory not found.")
        return actions

    docx_files = glob.glob(os.path.join(DATA_MINUTES, "**", "*.docx"), recursive=True)
    docx_files = [f for f in docx_files if not os.path.basename(f).startswith("~")]

    for filepath in sorted(docx_files):
        filename = os.path.basename(filepath)
        date_str = _parse_date_from_filename(filename)
        if not date_str:
            print(f"  [minutes] Could not parse date from: {filename}")
            continue

        print(f"  [minutes] Parsing: {filename} -> {date_str}")
        data = _parse_minutes_docx(filepath)

        # Convert action items to actions log format
        for item in data["action_items"]:
            actions.append({
                "date": date_str,
                "source": "Meeting Minutes",
                "source_party": "All Parties",
                "source_file": filename,
                "action": item["action"],
                "responsible": item["responsible"],
                "status": item["status"],
                "category": item["category"],
            })

        # Outlook entries
        for item in data["positive_outlook"]:
            actions.append({
                "date": date_str,
                "source": "Meeting Minutes",
                "source_party": "All Parties",
                "source_file": filename,
                "action": f"[Positive] {item}",
                "responsible": "",
                "status": "noted",
                "category": "outlook",
            })

        for item in data["negative_outlook"]:
            actions.append({
                "date": date_str,
                "source": "Meeting Minutes",
                "source_party": "All Parties",
                "source_file": filename,
                "action": f"[Concern] {item}",
                "responsible": "",
                "status": "monitoring",
                "category": "outlook",
            })

    print(f"  [minutes] Total action entries: {len(actions)}")
    return actions
