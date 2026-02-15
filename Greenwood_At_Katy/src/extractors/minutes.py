"""Extract action items and decisions from meeting minutes."""
import os
import re
import glob
from docx import Document
from src.config import DATA_MINUTES


def _parse_date_from_filename(filename):
    """Extract date from filename like '2026_02_06.docx'."""
    match = re.search(r'(\d{4})_(\d{2})_(\d{2})', filename)
    if match:
        return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
    return None


def _parse_minutes_docx(filepath):
    """Parse a meeting minutes DOCX and extract structured info."""
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also extract table data
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)

    result = {
        "summary": "",
        "key_metrics": [],
        "action_items": [],
        "positive_outlook": [],
        "negative_outlook": [],
        "full_text": "\n".join(paragraphs),
    }

    section = None
    for para in paragraphs:
        lower = para.lower()

        # Detect sections
        if "简介" in para:
            section = "intro"
        elif "运营状态" in para or "运营" in para and "状态" in para:
            section = "status"
        elif "最新进展" in para:
            section = "updates"
        elif "积极面" in para or "积极" in lower:
            section = "positive"
        elif "消极面" in para or "消极" in lower:
            section = "negative"
        elif "投资人决策" in para or "决策点" in para:
            section = "decisions"
        elif "需要关注" in para:
            section = "watchlist"
        elif "展望" in para:
            section = "outlook"

        # Extract based on section
        if section == "intro" and para != "简介" and not result["summary"]:
            result["summary"] = para

        elif section == "positive" and para != "积极面":
            if "：" in para or ":" in para:
                result["positive_outlook"].append(para)

        elif section == "negative" and para != "消极面":
            if "：" in para or ":" in para:
                result["negative_outlook"].append(para)

        elif section == "decisions":
            if ("批准" in para or "确认" in para or "决定" in para or
                    "需要" in para or "建立" in para or "输出" in para):
                # Determine responsible party
                responsible = "Unknown"
                if "Steven" in para or "steven" in lower:
                    responsible = "Steven (Pondmoon)"
                elif "Meredith" in para or "meredith" in lower:
                    responsible = "Meredith (AHC)"
                elif "Dantel" in para or "Danteil" in para:
                    responsible = "Danteil (Greystar)"

                result["action_items"].append({
                    "action": para.strip("- ").strip("[ ] Todo list").strip(),
                    "responsible": responsible,
                    "status": "pending",
                    "category": "decision",
                })

    # Extract key metrics from table
    for row in table_data:
        if len(row) >= 2 and row[0] and row[1]:
            result["key_metrics"].append({
                "metric": row[0],
                "value": row[1],
                "notes": row[2] if len(row) > 2 else "",
            })

    return result


def extract():
    """Extract action items from all meeting minutes files.

    Returns a list of action entries for the actions log.
    """
    actions = []

    if not os.path.exists(DATA_MINUTES):
        print("  [minutes] Data_Minutes directory not found.")
        return actions

    docx_files = glob.glob(os.path.join(DATA_MINUTES, "**", "*.docx"), recursive=True)

    for filepath in sorted(docx_files):
        filename = os.path.basename(filepath)
        date_str = _parse_date_from_filename(filename)
        if not date_str:
            print(f"  [minutes] Could not parse date from: {filename}")
            continue

        print(f"  [minutes] Parsing: {filename} -> {date_str}")
        data = _parse_minutes_docx(filepath)

        # Convert action items to actions log format
        for item in data["action_items"]:
            actions.append({
                "date": date_str,
                "source": "Meeting Minutes",
                "source_party": "All Parties",
                "source_file": filename,
                "action": item["action"],
                "responsible": item["responsible"],
                "status": item["status"],
                "category": item["category"],
            })

        # Also create entries from key outlook items
        for item in data["positive_outlook"]:
            actions.append({
                "date": date_str,
                "source": "Meeting Minutes",
                "source_party": "All Parties",
                "source_file": filename,
                "action": f"[Positive] {item}",
                "responsible": "",
                "status": "noted",
                "category": "outlook",
            })

        for item in data["negative_outlook"]:
            actions.append({
                "date": date_str,
                "source": "Meeting Minutes",
                "source_party": "All Parties",
                "source_file": filename,
                "action": f"[Concern] {item}",
                "responsible": "",
                "status": "monitoring",
                "category": "outlook",
            })

    print(f"  [minutes] Total action entries: {len(actions)}")
    return actions
